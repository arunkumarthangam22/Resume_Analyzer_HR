import io
import os
import re
import docx
import pdfplumber
import pytesseract
import spacy
from pdf2image import convert_from_path
from PIL import Image
from sentence_transformers import SentenceTransformer, util
import cloudinary.uploader # type: ignore
import requests
from dotenv import load_dotenv
from functools import lru_cache
from sentence_transformers import SentenceTransformer


# ✅ Load environment variables
load_dotenv()

# ✅ Poppler path for OCR in PDFs (change according to your system)
if os.name == "posix":  # Linux/macOS
    POPPLER_PATH = "/usr/bin"
else:  # Windows (local development)
    POPPLER_PATH = r"C:\Release-23.11.0-0\poppler-23.11.0\Library\bin"

# ✅ Lazy Load NLP & BERT models
@lru_cache(maxsize=1)
def get_spacy_model():
    return spacy.load("en_core_web_sm")


def get_bert_model():
    return SentenceTransformer("paraphrase-MiniLM-L3-v2", cache_folder="/tmp")


# ✅ Download Cloudinary File Before Processing
def download_file(file_url):
    """Downloads a file from Cloudinary for local processing."""
    try:
        response = requests.get(file_url)
        if response.status_code == 200:
            return io.BytesIO(response.content)
        else:
            print(f"❌ Failed to download file from Cloudinary: {file_url}")
            return None
    except Exception as e:
        print(f"❌ Error downloading file: {e}")
        return None

# ✅ Extract text from PDF (with OCR support)
def extract_text_from_pdf(file_obj):
    """Extracts text from a PDF file, including OCR for scanned resumes."""
    text = ""
    try:
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"

        if not text.strip():
            print("⚠️ No text extracted, attempting OCR...")
            text = extract_text_with_ocr(file_obj)
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
    
    return text.strip() if text else "❌ Error extracting text"

# ✅ OCR Function for Scanned PDFs (Optimized for Memory)
def extract_text_with_ocr(pdf_file):
    """Extracts text from scanned PDFs using OCR (optimized for memory)."""
    text = ""
    try:
        for image in convert_from_path(pdf_file, poppler_path=POPPLER_PATH):  # ✅ Process one image at a time
            text += pytesseract.image_to_string(image) + "\n"
            image.close()  # ✅ Free memory after processing
        print("✅ OCR successful!")
    except Exception as e:
        print(f"❌ OCR failed: {e}")
    return text.strip()

# ✅ Extract text from DOCX (including OCR for images)
def extract_text_from_docx(file_obj):
    """Extracts text from a DOCX file, including tables & images (via OCR)."""
    try:
        doc = docx.Document(file_obj)
        text = [para.text for para in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_text))

        full_text = "\n".join(text)

        if not full_text.strip():
            print("⚠️ No text found, trying OCR on embedded images...")
            full_text = extract_text_from_docx_images(file_obj)

        return full_text if full_text.strip() else "❌ No extractable text found."

    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return "❌ Error extracting text."

# ✅ OCR function for extracting text from images inside DOCX
def extract_text_from_docx_images(docx_path):
    """Extracts text from images embedded inside a DOCX file using OCR."""
    try:
        doc = docx.Document(docx_path)
        text = []

        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_part = doc.part.rels[rel].target_part
                image_bytes = io.BytesIO(image_part.blob)

                img = Image.open(image_bytes)
                extracted_text = pytesseract.image_to_string(img)
                text.append(extracted_text)

        return "\n".join(text) if text else "❌ No images found for OCR."
    except Exception as e:
        print(f"❌ OCR extraction failed: {e}")
        return "❌ OCR failed."

# ✅ Detect file type and extract text
def extract_text(file_obj):
    """Detects file type and extracts text from an uploaded file (PDF or DOCX)."""
    file_name = file_obj.name.lower()
    
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(file_obj)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_obj)
    else:
        return "❌ Unsupported file format"

# ✅ Extract Email and Phone Numbers from Text
def extract_email_and_phone(text):
    """Extracts email and phone number from resume text using regex."""
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    phone_pattern = r"(\+?\d{1,3}[-.\s]?)?(\(?\d{2,4}\)?[-.\s]?)?(\d{3,4}[-.\s]?\d{4})"

    email_matches = re.findall(email_pattern, text, re.IGNORECASE)
    phone_matches = re.findall(phone_pattern, text)

    email = email_matches[0] if email_matches else None
    phone_number = "".join(phone_matches[0]).replace(" ", "").replace("-", "").replace(".", "").replace("(", "").replace(")", "") if phone_matches else None

    return email, phone_number

# ✅ Optimized BERT-based similarity matching
def bert_match_keywords(bert_model, resume_text, job_text):
    """Compares resume text with job description using preloaded BERT model."""
    embeddings = bert_model.encode([resume_text, job_text], convert_to_tensor=True)
    similarity_score = util.pytorch_cos_sim(embeddings[0], embeddings[1]).item()
    return round(similarity_score * 100, 2)

# ✅ Optimized Multi-Factor ATS Scoring Function
def compute_ats_score(resume_text, job_data):
    """Compute ATS Score using lazy-loaded BERT model."""
    bert_model = get_bert_model()  # ✅ Load BERT model once

    scores = {
        "job_title": bert_match_keywords(bert_model, resume_text, job_data["job_title"]),
        "job_description": bert_match_keywords(bert_model, resume_text, job_data["job_description"]),
        "required_skills": bert_match_keywords(bert_model, resume_text, " ".join(job_data["required_skills"])),
        "preferred_qualifications": bert_match_keywords(bert_model, resume_text, " ".join(job_data["preferred_qualifications"])),
        "responsibilities": bert_match_keywords(bert_model, resume_text, " ".join(job_data["responsibilities"]))
    }

    weights = {"job_title": 20, "job_description": 30, "required_skills": 25, "preferred_qualifications": 10, "responsibilities": 15}
    final_ats_score = sum(scores[key] * weights[key] / 100 for key in scores) + 30


    # Add points based on ATS score range
    if final_ats_score < 30:
        final_ats_score += 30
    elif final_ats_score <= 40:
        final_ats_score += 35
    elif final_ats_score <= 50:
        final_ats_score += 40
    elif final_ats_score <= 70:
        final_ats_score += 25
    elif final_ats_score <= 80:
        final_ats_score += 10
    elif final_ats_score <= 90:
        final_ats_score += 5

        
    return {"scores": scores, "final_ats_score": round(final_ats_score, 2)}





